"""DOCX and HTML export mixin for ExportService."""

import html
import logging
from io import BytesIO

from src.memory.story_state import StoryState
from src.utils.validation import validate_not_none, validate_type

from ._base import ExportOptions, ExportServiceBase

logger = logging.getLogger(__name__)


class DocxExportMixin(ExportServiceBase):
    """Mixin providing DOCX and HTML export."""

    def to_html(
        self,
        state: StoryState,
        template: str | None = None,
        options: ExportOptions | None = None,
    ) -> str:
        """Export story as standalone HTML.

        Args:
            state: The story state to export.
            template: Template name to use. If None, uses ebook template.
            options: Custom export options. If None, uses template defaults.

        Returns:
            HTML formatted string.
        """
        logger.debug(f"Exporting story to HTML: {state.id}")

        # Get template and options
        tmpl = self.get_template(template)
        opts = options or tmpl.options

        brief = state.brief
        title = state.project_name or (brief.premise[:50] if brief else "Untitled Story")

        # Build default CSS with template options
        default_css = f"""
            body {{
                font-family: {opts.font_family};
                font-size: {opts.font_size}px;
                line-height: {opts.line_height};
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
            }}
            h1 {{
                color: #333;
                border-bottom: 2px solid #333;
                padding-bottom: 10px;
            }}
            h2 {{
                color: #555;
                margin-top: {opts.chapter_spacing}em;
                margin-bottom: 1em;
            }}
            .meta {{
                color: #666;
                font-style: italic;
                margin-bottom: 30px;
            }}
            p {{
                margin-bottom: {opts.paragraph_spacing}em;
            }}
        """

        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>{html.escape(title)}</title>",
            "<meta charset='utf-8'>",
            "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            "<style>",
            default_css,
            opts.custom_css,  # Add custom CSS from template
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{html.escape(title)}</h1>",
        ]

        if brief:
            html_parts.append(
                f"<p class='meta'>Genre: {html.escape(brief.genre)} | Tone: {html.escape(brief.tone)}<br>"
                f"Setting: {html.escape(brief.setting_place)}, {html.escape(brief.setting_time)}</p>"
            )

        for chapter in state.chapters:
            if chapter.content:
                chapter_header = self._format_chapter_header(chapter.number, chapter.title, opts)
                html_parts.append(f"<h2>{html.escape(chapter_header)}</h2>")

                for para in chapter.content.split("\n\n"):
                    if para.strip():
                        safe_para = html.escape(para.strip())
                        html_parts.append(f"<p>{safe_para}</p>")

        html_parts.extend(["</body>", "</html>"])
        return "\n".join(html_parts)

    def to_docx(
        self,
        state: StoryState,
        template: str | None = None,
        options: ExportOptions | None = None,
    ) -> bytes:
        """Export story as DOCX (Microsoft Word).

        Args:
            state: The story state to export.
            template: Template name to use. If None, uses ebook template.
            options: Custom export options. If None, uses template defaults.

        Returns:
            DOCX file as bytes.
        """
        validate_not_none(state, "state")
        validate_type(state, "state", StoryState)
        logger.info(f"Exporting story to DOCX: {state.id}")
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        # Get template and options
        tmpl = self.get_template(template)
        opts = options or tmpl.options

        doc = Document()

        # Set document margins from options
        sections = doc.sections
        for section in sections:
            margin = Inches(opts.page_margin_inches)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # Title
        brief = state.brief
        title = state.project_name or (brief.premise[:50] if brief else "Untitled Story")
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if brief:
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(
                f"Genre: {brief.genre} | Tone: {brief.tone}\n"
                f"Setting: {brief.setting_place}, {brief.setting_time}"
            )
            meta_run.font.size = Pt(10)
            meta_run.italic = True
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Map font family to DOCX font name
        font_map = {
            "Courier New, monospace": "Courier New",
            "Georgia, serif": "Georgia",
        }
        font_name = font_map.get(opts.font_family, "Georgia")

        # Chapters
        for chapter in state.chapters:
            if chapter.content:
                # Chapter heading
                chapter_header = self._format_chapter_header(chapter.number, chapter.title, opts)
                chapter_heading = doc.add_paragraph()
                chapter_run = chapter_heading.add_run(chapter_header)
                chapter_run.font.size = Pt(18)
                chapter_run.bold = True

                # Chapter content
                for para in chapter.content.split("\n\n"):
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_format = p.paragraph_format

                        # Apply font settings
                        for run in p.runs:
                            run.font.name = font_name
                            run.font.size = Pt(opts.font_size)

                        # Apply spacing settings
                        p_format.space_after = Pt(opts.paragraph_spacing * opts.font_size)
                        p_format.line_spacing = opts.line_height if not opts.double_spaced else 2.0

                # Page break after each chapter
                doc.add_page_break()

        # Write to bytes using context manager for proper cleanup
        with BytesIO() as output:
            doc.save(output)
            logger.debug("DOCX export complete")
            return output.getvalue()
