"""Export service - handles exporting stories to various formats."""

import html
import logging
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from memory.story_state import StoryState
from settings import STORIES_DIR, Settings
from utils.constants import get_language_code
from utils.validation import (
    validate_not_none,
    validate_string_in_choices,
    validate_type,
)

logger = logging.getLogger(__name__)


@dataclass
class ExportOptions:
    """Formatting options for exports."""

    # Font settings
    font_family: str = "Georgia, serif"
    font_size: int = 12  # in points for PDF/DOCX, pixels for HTML/EPUB
    line_height: float = 1.6

    # Spacing settings
    paragraph_spacing: float = 1.0  # em units
    chapter_spacing: float = 2.0  # em units

    # Chapter header settings
    chapter_number_format: str = "Chapter {number}"  # e.g., "Chapter 1" or "Ch. 1"
    include_chapter_numbers: bool = True
    chapter_separator: str = ": "  # separator between number and title

    # Custom CSS (for HTML/EPUB)
    custom_css: str = ""

    # Page settings (for PDF/DOCX)
    page_margin_inches: float = 1.0
    double_spaced: bool = False  # Override line_height to 2.0 if True


@dataclass
class ExportTemplate:
    """Template configuration for exports."""

    name: str
    description: str
    options: ExportOptions


# Built-in export templates
MANUSCRIPT_TEMPLATE = ExportTemplate(
    name="manuscript",
    description="Professional manuscript format - Courier New, 12pt, double-spaced, 1-inch margins",
    options=ExportOptions(
        font_family="Courier New, monospace",
        font_size=12,
        line_height=2.0,
        double_spaced=True,
        paragraph_spacing=0.0,  # No extra spacing in manuscript format
        chapter_spacing=3.0,
        page_margin_inches=1.0,
        chapter_number_format="Chapter {number}",
        include_chapter_numbers=True,
        chapter_separator=": ",
    ),
)

EBOOK_TEMPLATE = ExportTemplate(
    name="ebook",
    description="Reader-friendly ebook format - Georgia, readable spacing, clean headers",
    options=ExportOptions(
        font_family="Georgia, serif",
        font_size=14,
        line_height=1.8,
        double_spaced=False,
        paragraph_spacing=1.2,
        chapter_spacing=2.5,
        page_margin_inches=0.75,
        chapter_number_format="Chapter {number}",
        include_chapter_numbers=True,
        chapter_separator=": ",
        custom_css="""
            body { text-align: justify; }
            h2 { page-break-before: always; }
            p { margin-bottom: 1.2em; text-indent: 1.5em; }
            p:first-of-type { text-indent: 0; }
        """,
    ),
)

WEB_SERIAL_TEMPLATE = ExportTemplate(
    name="web_serial",
    description="Modern web serial format - clean, responsive, optimized for online reading",
    options=ExportOptions(
        font_family="-apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif",
        font_size=16,
        line_height=1.7,
        double_spaced=False,
        paragraph_spacing=1.0,
        chapter_spacing=2.0,
        page_margin_inches=1.0,
        chapter_number_format="Chapter {number}",
        include_chapter_numbers=True,
        chapter_separator=" - ",
        custom_css="""
            body {
                max-width: 650px;
                margin: 0 auto;
                padding: 20px;
                background: #fafafa;
                color: #333;
            }
            h1 {
                font-size: 2.5em;
                border-bottom: 3px solid #333;
                padding-bottom: 15px;
                margin-bottom: 30px;
            }
            h2 {
                font-size: 1.8em;
                margin-top: 50px;
                margin-bottom: 25px;
                color: #222;
            }
            .meta {
                color: #666;
                font-size: 0.9em;
                margin-bottom: 40px;
                padding: 15px;
                background: #f0f0f0;
                border-radius: 5px;
            }
            p {
                margin-bottom: 1em;
                line-height: 1.7;
            }
            @media (prefers-color-scheme: dark) {
                body { background: #1a1a1a; color: #e0e0e0; }
                h1 { border-bottom-color: #e0e0e0; }
                h2 { color: #f0f0f0; }
                .meta { background: #2a2a2a; color: #b0b0b0; }
            }
        """,
    ),
)

# Template registry
EXPORT_TEMPLATES: dict[str, ExportTemplate] = {
    "manuscript": MANUSCRIPT_TEMPLATE,
    "ebook": EBOOK_TEMPLATE,
    "web_serial": WEB_SERIAL_TEMPLATE,
}


def _validate_export_path(path: Path, base_dir: Path = STORIES_DIR.parent) -> Path:
    """Validate that an export path is safe to write to.

    Args:
        path: Path to validate
        base_dir: Base directory to constrain exports to (default: output/)

    Returns:
        Resolved absolute path

    Raises:
        ValueError: If path escapes base directory or is otherwise unsafe
    """
    try:
        resolved = path.resolve()
        base_resolved = base_dir.resolve()

        # Allow system temp directory for testing (cross-platform)
        temp_dir = Path(tempfile.gettempdir()).resolve()
        if resolved.is_relative_to(temp_dir):
            return resolved

        # Check if path is within base directory
        try:
            resolved.relative_to(base_resolved)
        except ValueError:
            raise ValueError(
                f"Export path {resolved} is outside allowed directory {base_resolved}"
            ) from None

        return resolved
    except (OSError, RuntimeError) as e:
        raise ValueError(f"Invalid export path: {path}") from e


class ExportService:
    """Export stories to various formats.

    Supports markdown, plain text, HTML, EPUB, PDF, and DOCX export with customizable templates.
    """

    def __init__(self, settings: Settings | None = None):
        """Initialize ExportService.

        Args:
            settings: Application settings. If None, loads from settings.json.
        """
        logger.debug("Initializing ExportService")
        self.settings = settings or Settings.load()
        logger.debug("ExportService initialized successfully")

    def get_template(self, template_name: str | None = None) -> ExportTemplate:
        """Get export template by name.

        Args:
            template_name: Name of the template. If None, returns ebook template.

        Returns:
            ExportTemplate configuration.

        Raises:
            ValueError: If template name is not found.
        """
        if template_name is None:
            template_name = "ebook"
            logger.debug("No template specified, using default 'ebook' template")

        if template_name not in EXPORT_TEMPLATES:
            raise ValueError(
                f"Unknown template '{template_name}'. "
                f"Available templates: {list(EXPORT_TEMPLATES.keys())}"
            )

        logger.debug(f"Retrieved export template: {template_name}")
        return EXPORT_TEMPLATES[template_name]

    def _format_chapter_header(
        self, chapter_number: int, chapter_title: str, options: ExportOptions
    ) -> str:
        """Format chapter header according to options.

        Args:
            chapter_number: Chapter number.
            chapter_title: Chapter title.
            options: Export options.

        Returns:
            Formatted chapter header.
        """
        if not options.include_chapter_numbers:
            logger.debug(f"Formatting chapter header without number: {chapter_title}")
            return chapter_title

        number_part = options.chapter_number_format.format(number=chapter_number)
        header = f"{number_part}{options.chapter_separator}{chapter_title}"
        logger.debug(f"Formatted chapter header: {header}")
        return header

    def to_markdown(self, state: StoryState) -> str:
        """Export story as markdown.

        Args:
            state: The story state to export.

        Returns:
            Markdown formatted string.
        """
        validate_not_none(state, "state")
        validate_type(state, "state", StoryState)
        logger.debug(f"Exporting story to markdown: {state.id}")
        brief = state.brief
        md_parts = []

        # Title and metadata
        title = state.project_name or (brief.premise[:50] if brief else "Untitled Story")
        md_parts.append(f"# {title}\n")

        if brief:
            md_parts.append(f"*Genre: {brief.genre} | Tone: {brief.tone}*\n")
            md_parts.append(f"*Setting: {brief.setting_place}, {brief.setting_time}*\n")
            if brief.themes:
                md_parts.append(f"*Themes: {', '.join(brief.themes)}*\n")
            md_parts.append("\n---\n")

        # Chapters
        for chapter in state.chapters:
            if chapter.content:
                md_parts.append(f"\n## Chapter {chapter.number}: {chapter.title}\n\n")
                md_parts.append(chapter.content)
                md_parts.append("\n")

        return "\n".join(md_parts)

    def to_text(self, state: StoryState) -> str:
        """Export story as plain text.

        Args:
            state: The story state to export.

        Returns:
            Plain text formatted string.
        """
        validate_not_none(state, "state")
        validate_type(state, "state", StoryState)
        logger.debug(f"Exporting story to plain text: {state.id}")
        brief = state.brief
        text_parts = []

        # Title
        title = state.project_name or (brief.premise[:80] if brief else "Untitled Story")
        text_parts.append(title.upper())
        text_parts.append("=" * len(title))
        text_parts.append("")

        if brief:
            text_parts.append(f"Genre: {brief.genre} | Tone: {brief.tone}")
            text_parts.append(f"Setting: {brief.setting_place}, {brief.setting_time}")
            text_parts.append("")
            text_parts.append("-" * 60)
            text_parts.append("")

        # Chapters
        for chapter in state.chapters:
            if chapter.content:
                text_parts.append(f"CHAPTER {chapter.number}: {chapter.title.upper()}")
                text_parts.append("")
                text_parts.append(chapter.content)
                text_parts.append("")
                text_parts.append("-" * 40)
                text_parts.append("")

        return "\n".join(text_parts)

    def to_epub(
        self,
        state: StoryState,
        template: str | None = None,
        options: ExportOptions | None = None,
    ) -> bytes:
        """Export story as EPUB e-book.

        Args:
            state: The story state to export.
            template: Template name to use. If None, uses ebook template.
            options: Custom export options. If None, uses template defaults.

        Returns:
            EPUB file as bytes.
        """
        validate_not_none(state, "state")
        validate_type(state, "state", StoryState)
        logger.info(f"Exporting story to EPUB: {state.id}")
        from ebooklib import epub

        # Get template and options
        tmpl = self.get_template(template)
        opts = options or tmpl.options

        book = epub.EpubBook()

        # Metadata
        brief = state.brief
        title = state.project_name or (brief.premise[:50] if brief else "Untitled Story")
        lang_code = get_language_code(brief.language) if brief else "en"

        book.set_identifier(state.id)
        book.set_title(title)
        book.set_language(lang_code)

        if brief:
            book.add_metadata("DC", "description", brief.premise)
            book.add_metadata("DC", "subject", brief.genre)

        # Create custom CSS for EPUB
        custom_style = f"""
            @namespace epub "http://www.idpf.org/2007/ops";

            body {{
                font-family: {opts.font_family};
                font-size: {opts.font_size}px;
                line-height: {opts.line_height};
                margin: 1em;
            }}

            h1 {{
                margin-top: {opts.chapter_spacing}em;
                margin-bottom: 1em;
                font-size: 1.5em;
            }}

            p {{
                margin-bottom: {opts.paragraph_spacing}em;
                text-align: justify;
            }}

            {opts.custom_css}
        """

        # Add CSS
        css = epub.EpubItem(
            uid="style",
            file_name="style.css",
            media_type="text/css",
            content=custom_style,
        )
        book.add_item(css)

        # Create chapters
        chapters = []
        for ch in state.chapters:
            if ch.content:
                # Format chapter header with options
                chapter_header = self._format_chapter_header(ch.number, ch.title, opts)
                safe_header = html.escape(chapter_header)

                epub_chapter = epub.EpubHtml(
                    title=chapter_header,
                    file_name=f"chapter_{ch.number}.xhtml",
                    lang=lang_code,
                )

                # Convert to HTML with escaped content (XSS prevention)
                paragraphs = ch.content.split("\n\n")
                html_paragraphs = [
                    f"<p>{html.escape(para)}</p>" for para in paragraphs if para.strip()
                ]
                html_content = "\n".join(html_paragraphs)

                # Use simple HTML structure without XML declaration (ebooklib handles that)
                epub_chapter.content = f"""
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{safe_header}</title>
    <link rel="stylesheet" href="style.css" type="text/css"/>
</head>
<body>
    <h1>{safe_header}</h1>
    {html_content}
</body>
</html>
"""

                book.add_item(epub_chapter)
                chapters.append(epub_chapter)

        # Navigation
        book.toc = tuple(chapters)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Spine
        book.spine = ["nav"] + chapters

        # Write to bytes
        output = BytesIO()
        epub.write_epub(output, book)
        logger.debug(f"EPUB export complete: {len(chapters)} chapters")
        return output.getvalue()

    def to_pdf(
        self,
        state: StoryState,
        template: str | None = None,
        options: ExportOptions | None = None,
    ) -> bytes:
        """Export story as PDF.

        Args:
            state: The story state to export.
            template: Template name to use. If None, uses ebook template.
            options: Custom export options. If None, uses template defaults.

        Returns:
            PDF file as bytes.
        """
        validate_not_none(state, "state")
        validate_type(state, "state", StoryState)
        logger.info(f"Exporting story to PDF: {state.id}")
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Flowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer

        # Get template and options
        tmpl = self.get_template(template)
        opts = options or tmpl.options

        buffer = BytesIO()

        # Use margin from options
        margin = opts.page_margin_inches * inch
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=margin,
            leftMargin=margin,
            topMargin=margin,
            bottomMargin=margin,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            spaceAfter=30,
        )

        # Calculate line spacing (leading = font_size * line_height)
        leading = opts.font_size * opts.line_height
        if opts.double_spaced:
            leading = opts.font_size * 2.0

        chapter_style = ParagraphStyle(
            "ChapterTitle",
            parent=styles["Heading2"],
            fontSize=18,
            spaceAfter=20,
            spaceBefore=opts.chapter_spacing * opts.font_size,
        )

        # Map font family to ReportLab font
        font_map = {
            "Courier New, monospace": "Courier",
            "Georgia, serif": "Times-Roman",
        }
        font_name = font_map.get(opts.font_family, "Times-Roman")

        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=opts.font_size,
            leading=leading,
            spaceAfter=opts.paragraph_spacing * opts.font_size,
        )

        story_elements: list[Flowable] = []

        # Title page
        brief = state.brief
        title = state.project_name or (brief.premise[:50] if brief else "Untitled Story")
        story_elements.append(Paragraph(title, title_style))

        if brief:
            story_elements.append(
                Paragraph(f"<i>{brief.genre} | {brief.tone}</i>", styles["Normal"])
            )
            story_elements.append(Spacer(1, 0.5 * inch))

        story_elements.append(PageBreak())

        # Chapters
        for ch in state.chapters:
            if ch.content:
                chapter_header = self._format_chapter_header(ch.number, ch.title, opts)
                story_elements.append(Paragraph(html.escape(chapter_header), chapter_style))

                # Split into paragraphs
                for para in ch.content.split("\n\n"):
                    if para.strip():
                        # Escape special characters using standard library
                        safe_para = html.escape(para.strip())
                        story_elements.append(Paragraph(safe_para, body_style))

                story_elements.append(PageBreak())

        doc.build(story_elements)
        logger.debug("PDF export complete")
        return buffer.getvalue()

    def to_html(
        self,
        state: StoryState,
        template: str | None = None,
        options: ExportOptions | None = None,
    ) -> str:
        """Export story as standalone HTML.

        Args:
            state: The story state to export.
            template: Template name to use. If None, uses ebook template.
            options: Custom export options. If None, uses template defaults.

        Returns:
            HTML formatted string.
        """
        logger.debug(f"Exporting story to HTML: {state.id}")

        # Get template and options
        tmpl = self.get_template(template)
        opts = options or tmpl.options

        brief = state.brief
        title = state.project_name or (brief.premise[:50] if brief else "Untitled Story")

        # Build default CSS with template options
        default_css = f"""
            body {{
                font-family: {opts.font_family};
                font-size: {opts.font_size}px;
                line-height: {opts.line_height};
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
            }}
            h1 {{
                color: #333;
                border-bottom: 2px solid #333;
                padding-bottom: 10px;
            }}
            h2 {{
                color: #555;
                margin-top: {opts.chapter_spacing}em;
                margin-bottom: 1em;
            }}
            .meta {{
                color: #666;
                font-style: italic;
                margin-bottom: 30px;
            }}
            p {{
                margin-bottom: {opts.paragraph_spacing}em;
            }}
        """

        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>{html.escape(title)}</title>",
            "<meta charset='utf-8'>",
            "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            "<style>",
            default_css,
            opts.custom_css,  # Add custom CSS from template
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{html.escape(title)}</h1>",
        ]

        if brief:
            html_parts.append(
                f"<p class='meta'>Genre: {html.escape(brief.genre)} | Tone: {html.escape(brief.tone)}<br>"
                f"Setting: {html.escape(brief.setting_place)}, {html.escape(brief.setting_time)}</p>"
            )

        for chapter in state.chapters:
            if chapter.content:
                chapter_header = self._format_chapter_header(chapter.number, chapter.title, opts)
                html_parts.append(f"<h2>{html.escape(chapter_header)}</h2>")

                for para in chapter.content.split("\n\n"):
                    if para.strip():
                        safe_para = html.escape(para.strip())
                        html_parts.append(f"<p>{safe_para}</p>")

        html_parts.extend(["</body>", "</html>"])
        return "\n".join(html_parts)

    def to_docx(
        self,
        state: StoryState,
        template: str | None = None,
        options: ExportOptions | None = None,
    ) -> bytes:
        """Export story as DOCX (Microsoft Word).

        Args:
            state: The story state to export.
            template: Template name to use. If None, uses ebook template.
            options: Custom export options. If None, uses template defaults.

        Returns:
            DOCX file as bytes.
        """
        validate_not_none(state, "state")
        validate_type(state, "state", StoryState)
        logger.info(f"Exporting story to DOCX: {state.id}")
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        # Get template and options
        tmpl = self.get_template(template)
        opts = options or tmpl.options

        doc = Document()

        # Set document margins from options
        sections = doc.sections
        for section in sections:
            margin = Inches(opts.page_margin_inches)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # Title
        brief = state.brief
        title = state.project_name or (brief.premise[:50] if brief else "Untitled Story")
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if brief:
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(
                f"Genre: {brief.genre} | Tone: {brief.tone}\n"
                f"Setting: {brief.setting_place}, {brief.setting_time}"
            )
            meta_run.font.size = Pt(10)
            meta_run.italic = True
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Map font family to DOCX font name
        font_map = {
            "Courier New, monospace": "Courier New",
            "Georgia, serif": "Georgia",
        }
        font_name = font_map.get(opts.font_family, "Georgia")

        # Chapters
        for chapter in state.chapters:
            if chapter.content:
                # Chapter heading
                chapter_header = self._format_chapter_header(chapter.number, chapter.title, opts)
                chapter_heading = doc.add_paragraph()
                chapter_run = chapter_heading.add_run(chapter_header)
                chapter_run.font.size = Pt(18)
                chapter_run.bold = True

                # Chapter content
                for para in chapter.content.split("\n\n"):
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_format = p.paragraph_format

                        # Apply font settings
                        for run in p.runs:
                            run.font.name = font_name
                            run.font.size = Pt(opts.font_size)

                        # Apply spacing settings
                        p_format.space_after = Pt(opts.paragraph_spacing * opts.font_size)
                        p_format.line_spacing = opts.line_height if not opts.double_spaced else 2.0

                # Page break after each chapter
                doc.add_page_break()

        # Write to bytes using context manager for proper cleanup
        with BytesIO() as output:
            doc.save(output)
            logger.debug("DOCX export complete")
            return output.getvalue()

    def save_to_file(
        self,
        state: StoryState,
        format: str,
        filepath: str | Path,
        template: str | None = None,
        options: ExportOptions | None = None,
    ) -> Path:
        """Save exported story to a file.

        Args:
            state: The story state to export.
            format: Export format ('markdown', 'text', 'epub', 'pdf', 'html', 'docx').
            filepath: Output file path.
            template: Template name to use. If None, uses ebook template.
            options: Custom export options. If None, uses template defaults.

        Returns:
            Path where the file was saved.

        Raises:
            ValueError: If format is not supported or path is invalid.
        """
        validate_not_none(state, "state")
        validate_type(state, "state", StoryState)
        validate_string_in_choices(
            format,
            "format",
            ["markdown", "text", "epub", "pdf", "html", "docx"],
        )
        validate_not_none(filepath, "filepath")
        logger.debug(
            f"save_to_file called: story_id={state.id}, format={format}, filepath={filepath}"
        )
        filepath = Path(filepath)

        try:
            # Validate export path
            filepath = _validate_export_path(filepath)
            filepath.parent.mkdir(parents=True, exist_ok=True)

            if format == "markdown":
                text_content = self.to_markdown(state)
                filepath.write_text(text_content, encoding="utf-8")
            elif format == "text":
                text_content = self.to_text(state)
                filepath.write_text(text_content, encoding="utf-8")
            elif format == "html":
                text_content = self.to_html(state, template, options)
                filepath.write_text(text_content, encoding="utf-8")
            elif format == "epub":
                bytes_content = self.to_epub(state, template, options)
                filepath.write_bytes(bytes_content)
            elif format == "pdf":
                bytes_content = self.to_pdf(state, template, options)
                filepath.write_bytes(bytes_content)
            elif format == "docx":
                bytes_content = self.to_docx(state, template, options)
                filepath.write_bytes(bytes_content)
            else:
                error_msg = f"Unsupported export format: {format}"
                logger.error(f"save_to_file failed: {error_msg}")
                raise ValueError(error_msg)

            logger.info(f"Exported story to {filepath} ({format} format)")
            return filepath
        except ValueError:
            raise
        except Exception as e:
            logger.error(
                f"save_to_file failed for story {state.id} to {filepath}: {e}", exc_info=True
            )
            raise

    def get_file_extension(self, format: str) -> str:
        """Get file extension for export format.

        Args:
            format: Export format name.

        Returns:
            File extension including dot.
        """
        logger.debug(f"get_file_extension called: format={format}")
        extensions = {
            "markdown": ".md",
            "text": ".txt",
            "html": ".html",
            "epub": ".epub",
            "pdf": ".pdf",
            "docx": ".docx",
        }
        return extensions.get(format, ".txt")
